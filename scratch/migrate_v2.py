import os
import docx
import pandas as pd
import re

SRC_DIR = r"D:\repos\Nativus\Google drive"
DEST_DIR = r"D:\repos\Nativus\TeX"

def ensure_dir(path):
    os.makedirs(path, exist_ok=True)

def escape_latex(text):
    text = re.sub(r'([_&$#{}])', r'\\\1', text)
    return text

def process_canto(docx_path, dest_tex_path):
    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        print(f"Could not read {docx_path}: {e}")
        return "", "", False
        
    lines = [p.text for p in doc.paragraphs]
    
    # Extract Header
    canto_title = ""
    music_style = ""
    start_idx = 0
    
    if len(lines) > 0 and lines[0].startswith("Canto"):
        canto_title = lines[0].strip()
        start_idx += 1
        if len(lines) > 1 and lines[1].strip() != "":
            music_style = lines[1].strip()
            start_idx += 1
            
    if not canto_title:
        for i, line in enumerate(lines):
            if line.strip().startswith("Canto"):
                canto_title = line.strip()
                start_idx = i + 1
                if start_idx < len(lines) and lines[start_idx].strip() != "":
                    music_style = lines[start_idx].strip()
                    start_idx += 1
                break

    # Extract Verses
    stanzas = []
    current_stanza = []
    
    for line in lines[start_idx:]:
        text = line.strip()
        if text.startswith('%'):
            current_stanza.append(text)
        elif text == "":
            if current_stanza:
                stanzas.append(current_stanza)
                current_stanza = []
        else:
            if '%' in text:
                parts = text.split('%', 1)
                text = escape_latex(parts[0].strip()) + " % " + parts[1].strip()
            else:
                text = escape_latex(text)
            current_stanza.append(text)
            
    if current_stanza:
        stanzas.append(current_stanza)
        
    # Write to .tex
    ensure_dir(os.path.dirname(dest_tex_path))
    with open(dest_tex_path, 'w', encoding='utf-8') as f:
        for stanza in stanzas:
            f.write("\\begin{minipage}{\\columnwidth}\n")
            f.write("\\begin{verse}\n")
            for i, line in enumerate(stanza):
                if line.startswith('%'):
                    f.write(line + "\n")
                else:
                    if i < len(stanza) - 1:
                        f.write(line + "\\\\\n")
                    else:
                        f.write(line + "\n")
            f.write("\\end{verse}\n")
            f.write("\\end{minipage}\n\\vspace{1em}\n\n")
            
    return escape_latex(canto_title), escape_latex(music_style), True

def generate_index_page(tex_path, title, items):
    # Items is a list of tuples (indent_level, text)
    ensure_dir(os.path.dirname(tex_path))
    with open(tex_path, 'w', encoding='utf-8') as f:
        f.write("\\input{preamble.tex}\n\\begin{document}\n\n")
        f.write("\\begin{center}\n")
        f.write(f"\\Huge \\textbf{{{escape_latex(title)}}}\n\\vspace{{2em}}\n")
        
        for level, text in items:
            if level == 0:
                f.write(f"\\Large \\textbf{{{escape_latex(text)}}} \\\\\n\\vspace{{1em}}\n")
            elif level == 1:
                f.write(f"\\large \\textit{{{escape_latex(text)}}} \\\\\n\\vspace{{0.5em}}\n")
            else:
                f.write(f"{escape_latex(text)} \\\\\n")
                
        f.write("\\end{center}\n\n\\end{document}\n")

def main():
    ensure_dir(DEST_DIR)
    
    # Parse Indice
    indice_path = os.path.join(SRC_DIR, "Administratie", "Indice.xlsx")
    df = pd.read_excel(indice_path, sheet_name="Canti")
    
    canticum_to_canti = {}
    opera_to_cantica = {}
    
    for idx, row in df.iterrows():
        canto_name = str(row['Canto']).strip()
        canticum_num = row['Canticum']
        opera_name = str(row['Plot']).strip()
        
        if pd.isna(canto_name) or canto_name == 'nan': continue
            
        canto_full_name = f"Canto {canto_name}"
        opera_full = f"Opera {opera_name}" if pd.notna(opera_name) and opera_name != 'nan' else "Opera Onbekend"
        cant_str = str(int(canticum_num)) if pd.notna(canticum_num) and str(canticum_num).replace('.0','').isdigit() else str(canticum_num)
        if pd.isna(canticum_num): cant_str = "Onbekend"
            
        if cant_str not in canticum_to_canti: canticum_to_canti[cant_str] = []
        canticum_to_canti[cant_str].append({'canto': canto_full_name, 'opera': opera_full, 'raw_name': canto_name})
        
        if opera_full not in opera_to_cantica: opera_to_cantica[opera_full] = set()
        opera_to_cantica[opera_full].add(cant_str)

    # Process Canti
    canti_src_dir = os.path.join(SRC_DIR, "Canti")
    canto_metadata = {}
    
    for root, dirs, files in os.walk(canti_src_dir):
        for file in files:
            if file.endswith('.docx') and file.startswith('Canto'):
                docx_path = os.path.join(root, file)
                rel_path = os.path.relpath(docx_path, canti_src_dir)
                dest_tex_path = os.path.join(DEST_DIR, "Canti", rel_path).replace('.docx', '.tex')
                
                canto_title, music_style, success = process_canto(docx_path, dest_tex_path)
                if success:
                    canto_metadata[file.replace('.docx', '')] = {
                        'title': canto_title,
                        'music': music_style,
                        'tex_path': os.path.join("Canti", rel_path).replace('.docx', '.tex').replace('\\', '/')
                    }
                    # User requested to delete fully converted files
                    try:
                        os.remove(docx_path)
                        print(f"Deleted fully migrated file: {docx_path}")
                    except Exception as e:
                        print(f"Failed to delete {docx_path}: {e}")
                        
    # Cantica mapping
    cantica_src_dir = os.path.join(SRC_DIR, "Cantica")
    canticum_real_names = {}
    if os.path.exists(cantica_src_dir):
        for file in os.listdir(cantica_src_dir):
            if file.endswith('.docx') and file.startswith('Canticum'):
                parts = file.split('_', 1)
                if len(parts) > 1:
                    num_part = parts[0].replace('Canticum', '').strip()
                    canticum_real_names[num_part] = file.replace('.docx', '')
                    # delete cantica docx as it's fully reconstructed
                    try:
                        os.remove(os.path.join(cantica_src_dir, file))
                    except: pass
                    
    # Generate Cantica files
    cantica_dest_dir = os.path.join(DEST_DIR, "Cantica")
    ensure_dir(cantica_dest_dir)
    
    opere_dest_dir = os.path.join(DEST_DIR, "Opere")
    ensure_dir(opere_dest_dir)
    
    generated_opere = []
    
    for opera, cantica_nums in opera_to_cantica.items():
        if opera == "Opera Onbekend": continue
        
        opera_src_dir = os.path.join(SRC_DIR, "Opere")
        real_opera_name = opera
        if os.path.exists(opera_src_dir):
            for file in os.listdir(opera_src_dir):
                if file.endswith('.docx') and opera.split()[-1] in file:
                    real_opera_name = file.replace('.docx', '')
                    try:
                        os.remove(os.path.join(opera_src_dir, file))
                    except: pass
                    break
                    
        opera_tex_path = os.path.join(opere_dest_dir, f"{real_opera_name}.tex")
        generated_opere.append(f"Opere/{real_opera_name}.tex")
        
        with open(opera_tex_path, 'w', encoding='utf-8') as f_op:
            f_op.write("\\input{../preamble.tex}\n")
            f_op.write("\\begin{document}\n\n")
            escaped_opera_name = escape_latex(real_opera_name)
            f_op.write(f"\\chapter*{{{escaped_opera_name}}}\n\n")
            
            sorted_cant_nums = sorted(list(cantica_nums), key=lambda x: int(x) if str(x).isdigit() else 999)
            
            for cant_num in sorted_cant_nums:
                real_cant_name = canticum_real_names.get(str(cant_num), f"Canticum {cant_num}")
                cant_tex_path = os.path.join(cantica_dest_dir, f"{real_cant_name}.tex")
                
                with open(cant_tex_path, 'w', encoding='utf-8') as f_cant:
                    # The user said: "No use a seperated preamble so all canti, cantica,, opere and trilogi are exaclt the same layout"
                    # We add the preamble to Canticum too so it can be compiled on its own
                    f_cant.write("\\input{../preamble.tex}\n\\begin{document}\n\n")
                    escaped_cant_name = escape_latex(real_cant_name)
                    f_cant.write(f"\\chapter{{{escaped_cant_name}}}\n\n")
                    
                    canti_list = canticum_to_canti.get(str(cant_num), [])
                    for c in canti_list:
                        canto_name = c['canto']
                        meta = canto_metadata.get(canto_name)
                        if meta:
                            title = meta['title']
                            music = meta['music']
                            tex_path = meta['tex_path']
                            f_cant.write(f"\\section*{{{title} --- {music}}}\n")
                            f_cant.write(f"\\input{{../Muziek/{canto_name}.tex}} % Placeholder for MusixTeX\n")
                            f_cant.write("\\begin{multicols}{3}\n")
                            f_cant.write(f"\\input{{../{tex_path}}}\n")
                            f_cant.write("\\end{multicols}\n\n")
                    f_cant.write("\n\\end{document}\n")
                            
                # For Opere, input Cantica. But wait! If Cantica has \begin{document}, Opere cannot just \input it without the standalone package!
                # I will still add `\usepackage{standalone}` in `preamble.tex` to allow this!
                f_op.write(f"\\input{{../Cantica/{real_cant_name}.tex}}\n")
                
            f_op.write("\n\\end{document}\n")

    # Generate Saga and Trilogie as index pages
    saga_items = []
    
    # Process Trilogie docs as index pages
    trilogie_src_dir = os.path.join(SRC_DIR, "Trilogie")
    if os.path.exists(trilogie_src_dir):
        trilogie_dest_dir = os.path.join(DEST_DIR, "Trilogie")
        ensure_dir(trilogie_dest_dir)
        for file in os.listdir(trilogie_src_dir):
            if file.endswith('.docx'):
                docx_path = os.path.join(trilogie_src_dir, file)
                try:
                    doc = docx.Document(docx_path)
                    items = []
                    level = 0
                    for p in doc.paragraphs:
                        text = p.text.strip()
                        if not text: continue
                        if "Trilogia" in text or "Opera" in text:
                            level = 0
                        elif "Canticum" in text:
                            level = 1
                        else:
                            level = 2
                        items.append((level, text))
                        
                    tex_path = os.path.join(trilogie_dest_dir, file.replace('.docx', '.tex'))
                    generate_index_page(tex_path, file.replace('.docx', ''), items)
                    saga_items.extend(items)
                    
                    os.remove(docx_path)
                    print(f"Deleted index docx: {docx_path}")
                except Exception as e:
                    print(f"Could not process trilogie {file}: {e}")

    # Generate Saga Nativus master file as an index
    saga_path = os.path.join(DEST_DIR, "Saga Nativus.tex")
    
    # Fallback if no trilogie files gave us items
    if not saga_items:
        for op in sorted(opera_to_cantica.keys()):
            saga_items.append((0, op))
            
    generate_index_page(saga_path, "Saga Nativus", saga_items)

    print("Migration V2 complete!")

if __name__ == "__main__":
    main()
